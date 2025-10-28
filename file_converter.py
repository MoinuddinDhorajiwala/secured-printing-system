"""
File conversion utilities for different document and image formats
Converts various formats to PDF for printing
"""

import os
import tempfile
from PIL import Image
import fitz  # PyMuPDF
import subprocess
import sys

def convert_to_pdf(input_file, output_file=None):
    """
    Convert various file formats to PDF
    Supported formats: PDF, DOCX, DOC, images (JPG, PNG, GIF, BMP, TIFF, WebP)
    """
    if not output_file:
        output_file = os.path.splitext(input_file)[0] + ".pdf"
    
    file_ext = os.path.splitext(input_file)[1].lower()
    
    try:
        if file_ext == '.pdf':
            # Already PDF, just copy
            import shutil
            shutil.copy2(input_file, output_file)
            return output_file
            
        elif file_ext in ['.docx', '.doc']:
            # Convert Word documents to PDF
            return convert_word_to_pdf(input_file, output_file)
            
        elif file_ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp']:
            # Convert images to PDF
            return convert_image_to_pdf(input_file, output_file)
            
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
            
    except Exception as e:
        raise Exception(f"Failed to convert {input_file} to PDF: {str(e)}")

def convert_word_to_pdf(input_file, output_file):
    """Convert Word documents to PDF using multiple advanced techniques"""
    try:
        # Technique 1: Try using LibreOffice first (most reliable)
        libreoffice_available = False
        try:
            result = subprocess.run([
                'libreoffice', '--version'
            ], capture_output=True, text=True, timeout=5)
            if result.returncode == 0:
                libreoffice_available = True
        except (subprocess.TimeoutExpired, FileNotFoundError):
            pass
        
        if libreoffice_available:
            try:
                result = subprocess.run([
                    'libreoffice', '--headless', '--convert-to', 'pdf',
                    '--outdir', os.path.dirname(output_file),
                    input_file
                ], capture_output=True, text=True, timeout=30)
                
                if result.returncode == 0:
                    expected_output = os.path.splitext(input_file)[0] + ".pdf"
                    if os.path.exists(expected_output):
                        import shutil
                        shutil.move(expected_output, output_file)
                        return output_file
                else:
                    print(f"LibreOffice conversion failed: {result.stderr}")
            except Exception as e:
                print(f"LibreOffice error: {e}")
        
        # Technique 2: Robust python-docx with text extraction
        if input_file.endswith('.docx'):
            try:
                from docx import Document
                import fitz
                
                # Load document
                doc = Document(input_file)
                
                # Create PDF with robust text extraction
                pdf_doc = fitz.open()
                page = pdf_doc.new_page()
                
                # Configuration
                margin_left = 50
                margin_top = 50
                line_height = 15
                max_width = 500
                
                y_position = margin_top
                
                def safe_insert_text(x, y, text, **kwargs):
                    """Safely insert text with fallback options"""
                    try:
                        # Try with default font
                        return page.insert_text((x, y), text, **kwargs)
                    except:
                        try:
                            # Try with Helvetica
                            return page.insert_text((x, y), text, fontname="Helvetica", **kwargs)
                        except:
                            # Ultimate fallback - use basic text
                            try:
                                return page.insert_text((x, y), text, fontname="helv", fontsize=12)
                            except:
                                # If all else fails, just return 0
                                return 0
                
                def process_paragraph(paragraph):
                    nonlocal y_position, page
                    
                    if not paragraph.text.strip():
                        return
                    
                    # Get paragraph style
                    style_name = str(paragraph.style.name) if hasattr(paragraph.style, 'name') else 'Normal'
                    
                    # Determine formatting based on style
                    font_size = 12
                    is_bold = False
                    is_heading = False
                    
                    if 'Title' in style_name:
                        font_size = 18
                        is_bold = True
                    elif 'Heading 1' in style_name:
                        font_size = 16
                        is_bold = True
                        is_heading = True
                    elif 'Heading 2' in style_name:
                        font_size = 14
                        is_bold = True
                        is_heading = True
                    elif 'Heading 3' in style_name:
                        font_size = 13
                        is_bold = True
                        is_heading = True
                    
                    # Add extra space before headings
                    if is_heading and y_position > margin_top + 20:
                        y_position += 10
                    
                    # Process text with word wrapping
                    text = paragraph.text
                    words = text.split()
                    current_line = ""
                    x_position = margin_left
                    
                    for word in words:
                        test_line = current_line + (" " if current_line else "") + word
                        estimated_width = len(test_line) * font_size * 0.6
                        
                        if estimated_width < max_width:
                            current_line = test_line
                        else:
                            # Draw current line
                            if current_line:
                                safe_insert_text(x_position, y_position, current_line, 
                                               fontsize=font_size, 
                                               fontname="helv-b" if is_bold else "helv")
                                y_position += line_height
                                
                                # Check for page break
                                if y_position > 750:
                                    page = pdf_doc.new_page()
                                    y_position = margin_top
                            
                            current_line = word
                    
                    # Draw remaining line
                    if current_line:
                        safe_insert_text(x_position, y_position, current_line, 
                                       fontsize=font_size, 
                                       fontname="helv-b" if is_bold else "helv")
                        y_position += line_height
                        
                        if y_position > 750:
                            page = pdf_doc.new_page()
                            y_position = margin_top
                    
                    # Add paragraph spacing
                    y_position += 5
                
                # Process all paragraphs
                for paragraph in doc.paragraphs:
                    process_paragraph(paragraph)
                
                # Process tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = ""
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text += cell.text.strip() + " | "
                        
                        if row_text:
                            safe_insert_text(margin_left, y_position, row_text.rstrip(" | "), 
                                           fontsize=10, fontname="helv")
                            y_position += 12
                            
                            if y_position > 750:
                                page = pdf_doc.new_page()
                                y_position = margin_top
                
                pdf_doc.save(output_file)
                pdf_doc.close()
                return output_file
                
            except Exception as e:
                print(f"Robust python-docx conversion failed: {e}")
        
        # Technique 3: Ultimate fallback - basic text extraction
        if input_file.endswith('.docx'):
            try:
                from docx import Document
                
                # Simple text extraction
                doc = Document(input_file)
                
                # Extract all text
                all_text = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        all_text.append(paragraph.text.strip())
                
                # Also extract table text
                for table in doc.tables:
                    for row in table.rows:
                        row_text = ""
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text += cell.text.strip() + " | "
                        if row_text:
                            all_text.append(row_text.rstrip(" | "))
                
                if not all_text:
                    raise Exception("No text content found in document")
                
                # Create simple PDF
                try:
                    import fitz
                    pdf_doc = fitz.open()
                    page = pdf_doc.new_page()
                    
                    y_position = 50
                    for text_line in all_text:
                        page.insert_text((50, y_position), text_line, fontsize=12, fontname="helv")
                        y_position += 15
                        
                        if y_position > 750:
                            page = pdf_doc.new_page()
                            y_position = 50
                    
                    pdf_doc.save(output_file)
                    pdf_doc.close()
                    
                except ImportError:
                    # Ultimate fallback: create text file and convert
                    text_file = output_file.replace('.pdf', '.txt')
                    with open(text_file, 'w', encoding='utf-8') as f:
                        f.write('\n'.join(all_text))
                    
                    # Try to convert text to PDF using reportlab if available
                    try:
                        from reportlab.lib.pagesizes import letter
                        from reportlab.pdfgen import canvas
                        from reportlab.lib.units import inch
                        
                        c = canvas.Canvas(output_file, pagesize=letter)
                        width, height = letter
                        
                        y_position = height - 1 * inch
                        for text_line in all_text:
                            c.drawString(1 * inch, y_position, text_line)
                            y_position -= 0.2 * inch
                            
                            if y_position < 1 * inch:
                                c.showPage()
                                y_position = height - 1 * inch
                        
                        c.save()
                        
                    except ImportError:
                        # Last resort: just return the text file
                        return text_file
                
                return output_file
                
            except Exception as e:
                raise Exception(f"Ultimate fallback conversion failed: {e}")
        
        if not libreoffice_available:
            raise Exception("LibreOffice not found. Please install LibreOffice for full Word document support, or use DOCX files with python-docx (enhanced conversion available).")
        else:
            raise Exception("All conversion methods failed. The document may be corrupted or in an unsupported format.")
        
    except Exception as e:
        raise Exception(f"Word to PDF conversion failed: {str(e)}")

def convert_image_to_pdf(input_file, output_file):
    """Convert images to PDF"""
    try:
        # Open image with PIL
        image = Image.open(input_file)
        
        # Convert to RGB if necessary (for JPEG compatibility)
        if image.mode != 'RGB':
            # Create white background for transparency
            background = Image.new('RGB', image.size, (255, 255, 255))
            if image.mode == 'RGBA':
                background.paste(image, mask=image.split()[3])
            else:
                background.paste(image)
            image = background
        
        # Save as PDF
        image.save(output_file, "PDF", resolution=100.0)
        
        return output_file
        
    except Exception as e:
        raise Exception(f"Image to PDF conversion failed: {str(e)}")

def is_conversion_required(file_path):
    """Check if file needs conversion to PDF"""
    ext = os.path.splitext(file_path)[1].lower()
    return ext not in ['.pdf']

def get_converted_pdf_path(original_file):
    """Get the path where the converted PDF would be saved"""
    base_name = os.path.splitext(original_file)[0]
    return base_name + "_converted.pdf"

# Test function
if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: python file_converter.py <input_file>")
        sys.exit(1)
    
    input_file = sys.argv[1]
    try:
        output_file = convert_to_pdf(input_file)
        print(f"Successfully converted {input_file} to {output_file}")
    except Exception as e:
        print(f"Conversion failed: {e}")
        sys.exit(1)